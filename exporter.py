from docx import Document as DocxDoc
from fpdf import FPDF

def exportar_a_word(preguntas, archivo_salida):
    doc = DocxDoc()
    for item in preguntas:
        doc.add_paragraph(f"Pregunta: {item['pregunta']}")
        doc.add_paragraph(f"Respuesta: {item['respuesta']}")
        doc.add_paragraph("")
    doc.save(archivo_salida)

def exportar_a_pdf(preguntas, archivo_salida):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for item in preguntas:
        pdf.multi_cell(0, 10, f"Pregunta: {item['pregunta']}\nRespuesta: {item['respuesta']}\n")
    pdf.output(archivo_salida)
